import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from src.logger import get_logger


    
from notebooklm import NotebookLMClient

logger = get_logger(__name__)


def _add_formatted_text(paragraph, text: str):
    """Hỗ trợ in đậm (**text**) và in nghiêng (*text*) khi ghi vào Word."""
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def _create_word_table(doc: Document, table_rows: list):
    """Tạo bảng Word chuẩn từ danh sách các dòng markdown phân tích được."""
    if not table_rows:
        return

    num_cols = len(table_rows[0])
    table = doc.add_table(rows=len(table_rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for row_idx, row_data in enumerate(table_rows):
        for col_idx, cell_value in enumerate(row_data):
            if col_idx < num_cols:
                cell = table.cell(row_idx, col_idx)
                p = cell.paragraphs[0]
                _add_formatted_text(p, cell_value.strip())
                if row_idx == 0:
                    for run in p.runs:
                        run.bold = True
    doc.add_paragraph()


def save_markdown_to_docx(markdown_text: str, output_path: str, title: str = "BÁO CÁO TỔNG HỢP"):
    """Phân tích Markdown và lưu thành tài liệu .docx hoàn chỉnh."""
    doc = Document()
    main_title = doc.add_heading(title, level=0)
    main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    lines = markdown_text.split("\n")
    table_buffer = []
    in_table = False

    for line in lines:
        stripped = line.strip()
        if stripped.startswith("|") and stripped.endswith("|"):
            if re.match(r'^\|[\s\-:|]+\|$', stripped):
                continue
            cols = [col.strip() for col in stripped.split("|")[1:-1]]
            table_buffer.append(cols)
            in_table = True
            continue
        else:
            if in_table:
                _create_word_table(doc, table_buffer)
                table_buffer = []
                in_table = False

        if not stripped:
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_text(p, stripped[2:])
        elif re.match(r'^\d+\.\s', stripped):
            content = re.sub(r'^\d+\.\s', '', stripped)
            p = doc.add_paragraph(style='List Number')
            _add_formatted_text(p, content)
        else:
            p = doc.add_paragraph()
            _add_formatted_text(p, stripped)

    if in_table and table_buffer:
        _create_word_table(doc, table_buffer)
        
    doc.save(output_path)
    
    
    

PROMPT_TEMPLATE = """
Hãy tóm tắt và phân tích chi tiết hồ sơ tài liệu này dưới dạng báo cáo chuyên nghiệp.
Yêu cầu định dạng cấu trúc chuẩn Markdown để chuyển đổi sang Word:

1. TIÊU ĐỀ: Sử dụng Tiêu đề cấp 1 (#), cấp 2 (##), cấp 3 (###) rõ ràng.
2. QUY TẮC TRÍCH DẪN (QUAN TRỌNG):
   - ƯU TIÊN SỐ TRANG IN ĐƯỢC ĐÁNH DƯỚI CHÂN TRANG VĂN BẢN (Page Number thực tế) và TÊN MỤC / CHƯƠNG / ĐIỀU KHOẢN (ví dụ: Chương III - Mục 2.1, Trang 32).
   - Nếu không chắc chắn số trang in, hãy trích dẫn kèm [Chương X, Mục Y, Bảng Z] để dễ tra cứu chính xác thay vì chỉ ghi số trang PDF.
   - Định dạng trích dẫn: (Trang X, Mục Y) hoặc (Chương X, Điều Y).
3. BẢNG BIỂU: Trình bày thông số, số lượng, tiêu chuẩn bằng bảng Markdown chuẩn (| Cột 1 | Cột 2 | Cột 3 |). Có cột riêng cho vị trí trích dẫn.
4. Bôi đậm các thông số quan trọng, mốc thời gian, số tiền (**nội dung**).
5. Trả về trực tiếp nội dung báo cáo, không thêm lời mở đầu hay kết bài.
"""


async def generate_report_docx(pdf_path: str, docx_path: str, ma_tbmt: str) -> bool:
    """
    Tải PDF lên NotebookLM tạm thời, tóm tắt Markdown và lưu ra file Word.
    Trả về True nếu thành công, False nếu có lỗi.
    """
    try:
        async with NotebookLMClient.from_storage() as client:
            nb = await client.notebooks.create(title=f"Temp_{ma_tbmt}")
            notebook_id = nb.id
            try:
                source = await client.sources.add_file(notebook_id, pdf_path)
                await client.sources.wait_for_sources(notebook_id, source_ids=[source.id])
                result = await client.chat.ask(notebook_id, PROMPT_TEMPLATE)
                
                save_markdown_to_docx(
                    markdown_text=result.answer,
                    output_path=docx_path,
                    title=f"BÁO CÁO PHÂN TÍCH HỒ SƠ: {ma_tbmt}"
                )
                return True
            finally:
                # Luôn giải phóng notebook tạm
                await client.notebooks.delete(notebook_id)
    except Exception as e:
        logger.error(f"Lỗi tạo DOCX qua NotebookLM cho gói {ma_tbmt}: {e}")
        return False