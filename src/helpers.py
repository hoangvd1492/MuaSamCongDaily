from datetime import datetime
from urllib.parse import urlencode
import random
import time
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls


import asyncio


async def async_sleep(min_ms=1000, max_ms=5000):
    await asyncio.sleep(random.randint(min_ms, max_ms) / 1000)


# Alias để tương thích
async def sleep(min_ms=1000, max_ms=5000):
    await async_sleep(min_ms, max_ms)


def query_builder(keyword, page, from_time, to_time):
    return [
        {
            "pageSize": 10,
            "pageNumber": page,
            "query": [
                {
                    "index": "es-contractor-selection",
                    "keyWord": keyword,
                    "matchType": "exact",
                    "matchFields": ["notifyNo", "bidName"],
                    "filters": [
                        {
                            "fieldName": "type",
                            "searchType": "in",
                            "fieldValues": ["es-notify-contractor"],
                        },
                        {
                            "fieldName": "investField",
                            "searchType": "in",
                            "fieldValues": ["PTV", "HH", "HON_HOP"],
                        },
                        {
                            "fieldName": "caseKHKQ",
                            "searchType": "not_in",
                            "fieldValues": ["1"],
                        },
                        {
                            "fieldName": "publicDate",
                            "searchType": "range",
                            "from": from_time,
                            "to": to_time,
                        },
                    ],
                }
            ],
        }
    ]




PROMPT_TEMPLATE = """
    Hãy đóng vai một chuyên gia phân tích hồ sơ thầu hàng đầu. Hãy đọc kỹ toàn bộ Hồ sơ mời thầu (HSMT) đính kèm và lập một BÁO CÁO PHÂN TÍCH HỒ SƠ MỜI THẦU dưới dạng MARKDOWN trình bày khoa học và chuyên nghiệp.

    Yêu cầu cấu trúc Báo cáo Markdown như sau (sử dụng bảng Table Markdown cho Nhân sự, Thiết bị, Tài chính):

    # 📋 BÁO CÁO PHÂN TÍCH HỒ SƠ MỜI THẦU

    ## I. THÔNG TIN CHUNG GÓI THẦU
    - **Tên gói thầu:** 
    - **Bên mời thầu / Chủ đầu tư:** 
    - **Thời gian thực hiện hợp đồng:** 
    - **Nguồn vốn:** 
    - **Hình thức lựa chọn nhà thầu & Phương thức đấu thầu:** 
    - **Thời điểm đóng thầu / Hiệu lực E-HSDT:** 

    ## II. YÊU CẦU NĂNG LỰC TÀI CHÍNH & KINH NGHIỆM
    | Tiêu chí | Yêu cầu chi tiết |
    | :--- | :--- |
    | **Bảo đảm dự thầu** | |
    | **Doanh thu bình quân hàng năm** | |
    | **Hợp đồng tương tự** | |
    | **Yêu cầu nguồn lực tài chính / Tín dụng** | |

    ## III. YÊU CẦU NHÂN SỰ CHỦ CHỐT
    Trình bày dưới dạng Bảng Markdown:
    | STT | Vị trí nhân sự | Số lượng | Trình độ bằng cấp & Chứng chỉ | Kinh nghiệm công tác |
    | :-: | :--- | :-: | :--- | :--- |

    ## IV. YÊU CẦU THIẾT BỊ / MÁY MÓC CHỦ YẾU (nếu có)
    Trình bày dưới dạng Bảng Markdown (Nếu không yêu cầu thì ghi rõ "Không yêu cầu"):
    | STT | Tên loại thiết bị / máy móc | Số lượng | Yêu cầu tính năng / Công suất |
    | :-: | :--- | :-: | :--- |

    ## V. CÁC ĐIỀU KIỆN & YÊU CẦU ĐẶC BIỆT KHÁC
    - **Loại hợp đồng & Đơn giá:** 
    - **Tỷ lệ thầu phụ tối đa:** 
    - **Bảo đảm thực hiện hợp đồng:** 
    - **Các yêu cầu tư cách hợp lệ / Chứng chỉ / Giấy phép đặc thù:** 
    - **Các lưu ý quan trọng khác:** 

    Chỉ trả về nội dung Báo cáo Markdown, không kèm theo câu mở đầu hay giải thích gì khác.
"""



PROCESS_APPLY_MAP = {
  'LDT': "Luật Đấu thầu/ Áp dụng Luật Đấu thầu",
  'ADB': "ADB (Qua mạng)",
  'WB': "WB (Qua mạng)",
  'CPTPP': "CPTPP",
  'EVFTA': "EVFTA/UKVFTA",
  'UKFTA': "CPTPP/EVFTA/UKVFTA",
  'KHAC': "Khác (bao gồm: Viettel/ ADB/WB không qua mạng/ Áp dụng các điều khoản điểm cụ thể của Luật đấu thầu)",
}

BID_FORM_MAP = {
  'CGTT': "Chào giá trực tuyến",
  'CGTTRG': "Chào giá trực tuyến theo quy trình rút gọn",
  'DTRR': "Đấu thầu rộng rãi",
  'DH_GNV': "Đặt hàng, giao nhiệm vụ",
  'ESHOP': "Mua sắm trực tuyến",
  'DH': "Đặt hàng",
  'GNV': "Giao nhiệm vụ",
  'CHCT': "Chào hàng cạnh tranh",
  'CHCTRG': "Chào hàng cạnh tranh rút gọn",
  'DTHC': "Đấu thầu hạn chế",
  'MSTT': "Mua sắm trực tiếp",
  'CDT': "Chỉ định thầu",
  'CDTRG': "Chỉ định thầu rút gọn",
  'TTH': "Tự thực hiện",
  'LCNT_DB': "Lựa chọn nhà thầu trong trường hợp đặc biệt",
  'TVCN': "Tư vấn cá nhân",
  'TCTVCN': "Tuyển chọn tư vấn cá nhân",
  'DPCT': "Đàm phán cạnh tranh",
  'QCBS': "Tuyển chọn trên cơ sở Chất lượng và Chi phí (QCBS)",
  'DPG': "Đàm phán giá",
  'QBS': "Tuyển chọn tư vấn dựa trên cơ sở chất lượng (QBS)",
  'FBS': "Lựa chọn theo mức ngân sách cố định (FBS)",
  'LCS': "Tuyển chọn tư vấn có chi phí thấp nhất (LCS)",
  'CQS': "Tuyển chọn dựa trên năng lực (CQS)",
  'SSS': "Tuyển chọn từ một nguồn duy nhất (SSS)",
  'TGTC': "Tuyển chọn tư vấn trong các khoản vay cho các thể chế hoặc tổ chức trung gian tài chính",
  'NHBD': "Tuyển chọn tư vấn trong các khoản vay được Ngân hàng bảo đảm",
  'TVCT': "Tuyển chọn một số loại hình tư vấn cụ thể",
  'TGTHCD': "Tham gia thực hiện cộng đồng",
  'TVCNRG': "Tư vấn cá nhân rút gọn",
  'KHAC': "Khác",
}

BID_MODE_MAP = {
  "1_MTHS": "Một giai đoạn một túi hồ sơ",
  "1_HTHS": "Một giai đoạn hai túi hồ sơ",
  "2_MTHS": "Hai giai đoạn một túi hồ sơ",
  "2_HTHS": "Hai giai đoạn hai túi hồ sơ",
}

INVEST_FIELD_MAP = {
  'HH': "Hàng hóa",
  'XL': "Xây lắp",
  'PTV': "Phi tư vấn",
  'TV': "Tư vấn",
  'HON_HOP': "Hỗn hợp",
}


CONTRACT_TYPE_MAP = {
  'TG': "Trọn gói",
  'TG_DGCD': "Trọn gói và đơn giá cố định",
  'TG_CD_DC': "Trọn gói, đơn giá cố định và đơn giá điều chỉnh",
  'TG_CD_DC_TTG':
    "Trọn gói, đơn giá cố định, đơn giá điều chỉnh và theo thời gian",
  'TG_CD_DC_TTG_CPCP':
    "Trọn gói, đơn giá cố định, đơn giá điều chỉnh, theo thời gian và theo chi phí cộng phí",
  'TG_CD_DC_TTG_CPCP_KQDR':
    "Trọn gói, đơn giá cố định, đơn giá điều chỉnh, theo thời gian, theo chi phí cộng phí và theo kết quả đầu ra",
  'TG_CD_DC_TTG_CPCP_KQDR_TLPT':
    "Trọn gói, đơn giá cố định, đơn giá điều chỉnh, theo thời gian, theo chi phí cộng phí, theo kết quả đầu ra và theo tỷ lệ phần trăm",
  'TG_CD_DC_TTG_CPCP_TLPT':
    "Trọn gói, đơn giá cố định, đơn giá điều chỉnh, theo thời gian, theo chi phí cộng phí và theo tỷ lệ phần trăm",
  'TG_CD_DC_TTG_KQDR':
    "Trọn gói, đơn giá cố định, đơn giá điều chỉnh, theo thời gian và theo kết quả đầu ra",
  'TG_CD_DC_TTG_KQDR_TLPT':
    "Trọn gói, đơn giá cố định, đơn giá điều chỉnh, theo thời gian, theo kết quả đầu ra và theo tỷ lệ phần trăm",
  'TG_CD_DC_TTG_TLPT':
    "Trọn gói, đơn giá cố định, đơn giá điều chỉnh, theo thời gian và theo tỷ lệ phần trăm",
  'TG_CD_DC_CPCP':
    "Trọn gói, đơn giá cố định, đơn giá điều chỉnh và theo chi phí cộng phí",
  'TG_CD_DC_CPCP_KQDR':
    "Trọn gói, đơn giá cố định, đơn giá điều chỉnh, theo chi phí cộng phí và theo kết quả đầu ra",
  'TG_CD_DC_CPCP_KQDR_TLPT':
    "Trọn gói, đơn giá cố định, đơn giá điều chỉnh, theo chi phí cộng phí, theo kết quả đầu ra và theo tỷ lệ phần trăm",
  'TG_CD_DC_CPCP_TLPT':
    "Trọn gói, đơn giá cố định, đơn giá điều chỉnh, theo chi phí cộng phí và theo tỷ lệ phần trăm",
  'TG_CD_DC_KQDR':
    "Trọn gói, đơn giá cố định, đơn giá điều chỉnh và theo kết quả đầu ra",
  'TG_CD_DC_KQDR_TLPT':
    "Trọn gói, đơn giá cố định, đơn giá điều chỉnh, theo kết quả đầu ra và theo tỷ lệ phần trăm",
  'TG_CD_DC_TLPT':
    "Trọn gói, đơn giá cố định, đơn giá điều chỉnh và theo tỷ lệ phần trăm",
  'TG_CD_TTG': "Trọn gói, đơn giá cố định và theo thời gian",
  'TG_CD_TTG_CPCP':
    "Trọn gói, đơn giá cố định, theo thời gian và theo chi phí cộng phí",
  'TG_CD_TTG_CPCP_KQDR':
    "Trọn gói, đơn giá cố định, theo thời gian, theo chi phí cộng phí và theo kết quả đầu ra",
  'TG_CD_TTG_CPCP_KQDR_TLPT':
    "Trọn gói, đơn giá cố định, theo thời gian, theo chi phí cộng phí, theo kết quả đầu ra và theo tỷ lệ phần trăm",
  'TG_CD_TTG_CPCP_TLPT':
    "Trọn gói, đơn giá cố định, theo thời gian, theo chi phí cộng phí và theo tỷ lệ phần trăm",
  'TG_CD_TTG_KQDR':
    "Trọn gói, đơn giá cố định, theo thời gian và theo kết quả đầu ra",
  'TG_CD_TTG_KQDR_TLPT':
    "Trọn gói, đơn giá cố định, theo thời gian, theo kết quả đầu ra và theo tỷ lệ phần trăm",
  'TG_CD_TTG_TLPT':
    "Trọn gói, đơn giá cố định, theo thời gian và theo tỷ lệ phần trăm",
  'TG_CD_CPCP': "Trọn gói, đơn giá cố định và theo chi phí cộng phí",
  'TG_CD_CPCP_KQDR':
    "Trọn gói, đơn giá cố định, theo chi phí cộng phí và theo kết quả đầu ra",
  'TG_CD_CPCP_KQDR_TLPT':
    "Trọn gói, đơn giá cố định, theo chi phí cộng phí, theo kết quả đầu ra và theo tỷ lệ phần trăm",
  'TG_CD_CPCP_TLPT':
    "Trọn gói, đơn giá cố định, theo chi phí cộng phí và theo tỷ lệ phần trăm",
  'TG_CD_KQDR': "Trọn gói, đơn giá cố định và theo kết quả đầu ra",
  'TG_CD_KQDR_TLPT':
    "Trọn gói, đơn giá cố định, theo kết quả đầu ra và theo tỷ lệ phần trăm",
  'TG_CD_TLPT': "Trọn gói, đơn giá cố định và theo tỷ lệ phần trăm",
  'TG_DGDC': "Trọn gói và đơn giá điều chỉnh",
  'TG_DC_TTG': "Trọn gói, đơn giá điều chỉnh và theo thời gian",
  'TG_DC_TTG_CPCP':
    "Trọn gói, đơn giá điều chỉnh, theo thời gian và theo chi phí cộng phí",
  'TG_DC_TTG_CPCP_KQDR':
    "Trọn gói, đơn giá điều chỉnh, theo thời gian, theo chi phí cộng phí và theo kết quả đầu ra",
  'TG_DC_TTG_CPCP_KQDR_TLPT':
    "Trọn gói, đơn giá điều chỉnh, theo thời gian, theo chi phí cộng phí, theo kết quả đầu ra và theo tỷ lệ phần trăm",
  'TG_DC_TTG_CPCP_TLPT':
    "Trọn gói, đơn giá điều chỉnh, theo thời gian, theo chi phí cộng phí và theo tỷ lệ phần trăm",
  'TG_DC_TTG_KQDR':
    "Trọn gói, đơn giá điều chỉnh, theo thời gian và theo kết quả đầu ra",
  'TG_DC_TTG_KQDR_TLPT':
    "Trọn gói, đơn giá điều chỉnh, theo thời gian, theo kết quả đầu ra và theo tỷ lệ phần trăm",
  'TG_DC_TTG_TLPT':
    "Trọn gói, đơn giá điều chỉnh, theo thời gian và theo tỷ lệ phần trăm",
  'TG_DC_CPCP': "Trọn gói, đơn giá điều chỉnh và theo chi phí cộng phí",
  'TG_DC_CPCP_KQDR':
    "Trọn gói, đơn giá điều chỉnh, theo chi phí cộng phí và theo kết quả đầu ra",
  'TG_DC_CPCP_KQDR_TLPT':
    "Trọn gói, đơn giá điều chỉnh, theo chi phí cộng phí, theo kết quả đầu ra và theo tỷ lệ phần trăm",
  'TG_DC_CPCP_TLPT':
    "Trọn gói, đơn giá điều chỉnh, theo chi phí cộng phí và theo tỷ lệ phần trăm",
  'TG_DC_KQDR': "Trọn gói, đơn giá điều chỉnh và theo kết quả đầu ra",
  'TG_DC_KQDR_TLPT':
    "Trọn gói, đơn giá điều chỉnh, theo kết quả đầu ra và theo tỷ lệ phần trăm",
  'TG_DC_TLPT': "Trọn gói, đơn giá điều chỉnh và theo tỷ lệ phần trăm",
  'TG_TTG': "Trọn gói và theo thời gian",
  'TG_TTG_CPCP': "Trọn gói, theo thời gian và theo chi phí cộng phí",
  'TG_TTG_CPCP_KQDR':
    "Trọn gói, theo thời gian, theo chi phí cộng phí và theo kết quả đầu ra",
  'TG_TTG_CPCP_KQDR_TLPT':
    "Trọn gói, theo thời gian, theo chi phí cộng phí, theo kết quả đầu ra và theo tỷ lệ phần trăm",
  'TG_TTG_CPCP_TLPT':
    "Trọn gói, theo thời gian, theo chi phí cộng phí và theo tỷ lệ phần trăm",
  'TG_TTG_KQDR': "Trọn gói, theo thời gian và theo kết quả đầu ra",
  'TG_TTG_KQDR_TLPT':
    "Trọn gói, theo thời gian, theo kết quả đầu ra và theo tỷ lệ phần trăm",
  'TG_TTG_TLPT': "Trọn gói, theo thời gian và theo tỷ lệ phần trăm",
  'TG_CPCP': "Trọn gói và theo chi phí cộng phí",
  'TG_CPCP_KQDR': "Trọn gói, theo chi phí cộng phí và theo kết quả đầu ra",
  'TG_CPCP_KQDR_TLPT':
    "Trọn gói, theo chi phí cộng phí, theo kết quả đầu ra và theo tỷ lệ phần trăm",
  'TG_CPCP_TLPT': "Trọn gói, theo chi phí cộng phí và theo tỷ lệ phần trăm",
  'TG_KQDR': "Trọn gói và theo kết quả đầu ra",
  'TG_KQDR_TLPT': "Trọn gói, theo kết quả đầu ra và theo tỷ lệ phần trăm",
  'TG_TLPT': "Trọn gói và theo tỷ lệ phần trăm",
  'DGCD': "Đơn giá cố định",
  'DGCD_DC': "Đơn giá cố định và đơn giá điều chỉnh",
  'CD_DC_TTG': "Đơn giá cố định, đơn giá điều chỉnh và theo thời gian",
  'CD_DC_TTG_CPCP':
    "Đơn giá cố định, đơn giá điều chỉnh, theo thời gian và theo chi phí cộng phí",
  'CD_DC_TTG_CPCP_KQDR':
    "Đơn giá cố định, đơn giá điều chỉnh, theo thời gian, theo chi phí cộng phí và theo kết quả đầu ra",
  'CD_DC_TTG_CPCP_KQDR_TLPT':
    "Đơn giá cố định, đơn giá điều chỉnh, theo thời gian, theo chi phí cộng phí, theo kết quả đầu ra và theo tỷ lệ phần trăm",
  'CD_DC_TTG_CPCP_TLPT':
    "Đơn giá cố định, đơn giá điều chỉnh, theo thời gian, theo chi phí cộng phí và theo tỷ lệ phần trăm",
  'CD_DC_TTG_KQDR':
    "Đơn giá cố định, đơn giá điều chỉnh, theo thời gian và theo kết quả đầu ra",
  'CD_DC_TTG_KQDR_TLPT':
    "Đơn giá cố định, đơn giá điều chỉnh, theo thời gian, theo kết quả đầu ra và theo tỷ lệ phần trăm",
  'CD_DC_TTG_TLPT':
    "Đơn giá cố định, đơn giá điều chỉnh, theo thời gian và theo tỷ lệ phần trăm",
  'CD_DC_CPCP': "Đơn giá cố định, đơn giá điều chỉnh và theo chi phí cộng phí",
  'CD_DC_CPCP_KQDR':
    "Đơn giá cố định, đơn giá điều chỉnh, theo chi phí cộng phí và theo kết quả đầu ra",
  'CD_DC_CPCP_KQDR_TLPT':
    "Đơn giá cố định, đơn giá điều chỉnh, theo chi phí cộng phí, theo kết quả đầu ra và theo tỷ lệ phần trăm",
  'CD_DC_CPCP_TLPT':
    "Đơn giá cố định, đơn giá điều chỉnh, theo chi phí cộng phí và theo tỷ lệ phần trăm",
  'CD_DC_KQDR': "Đơn giá cố định, đơn giá điều chỉnh và theo kết quả đầu ra",
  'CD_DC_KQDR_TLPT':
    "Đơn giá cố định, đơn giá điều chỉnh, theo kết quả đầu ra và theo tỷ lệ phần trăm",
  'CD_DC_TLPT': "Đơn giá cố định, đơn giá điều chỉnh và theo tỷ lệ phần trăm",
  'CD_TTG': "Đơn giá cố định và theo thời gian",
  'CD_TTG_CPCP': "Đơn giá cố định, theo thời gian và theo chi phí cộng phí",
  'CD_TTG_CPCP_KQDR':
    "Đơn giá cố định, theo thời gian, theo chi phí cộng phí và theo kết quả đầu ra",
  'CD_TTG_CPCP_KQDR_TLPT':
    "Đơn giá cố định, theo thời gian, theo chi phí cộng phí, theo kết quả đầu ra và theo tỷ lệ phần trăm",
  'CD_TTG_CPCP_TLPT':
    "Đơn giá cố định, theo thời gian, theo chi phí cộng phí và theo tỷ lệ phần trăm",
  'CD_TTG_KQDR': "Đơn giá cố định, theo thời gian và theo kết quả đầu ra",
  'CD_TTG_KQDR_TLPT':
    "Đơn giá cố định, theo thời gian, theo kết quả đầu ra và theo tỷ lệ phần trăm",
  'CD_TTG_TLPT': "Đơn giá cố định, theo thời gian và theo tỷ lệ phần trăm",
  'CD_CPCP': "Đơn giá cố định và theo chi phí cộng phí",
  'CD_CPCP_KQDR': "Đơn giá cố định, theo chi phí cộng phí và theo kết quả đầu ra",
  'CD_CPCP_KQDR_TLPT':
    "Đơn giá cố định, theo chi phí cộng phí, theo kết quả đầu ra và theo tỷ lệ phần trăm",
  'CD_CPCP_TLPT':
    "Đơn giá cố định, theo chi phí cộng phí và theo tỷ lệ phần trăm",
  'CD_KQDR': "Đơn giá cố định và theo kết quả đầu ra",
  'CD_KQDR_TLPT': "Đơn giá cố định, theo kết quả đầu ra và theo tỷ lệ phần trăm",
  'CD_TLPT': "Đơn giá cố định và theo tỷ lệ phần trăm",
  'DGDC': "Đơn giá điều chỉnh",
  'DC_TTG': "Đơn giá điều chỉnh và theo thời gian",
  'DC_TTG_CPCP': "Đơn giá điều chỉnh, theo thời gian và theo chi phí cộng phí",
  'DC_TTG_CPCP_KQDR':
    "Đơn giá điều chỉnh, theo thời gian, theo chi phí cộng phí và theo kết quả đầu ra",
  'DC_TTG_CPCP_KQDR_TLPT':
    "Đơn giá điều chỉnh, theo thời gian, theo chi phí cộng phí, theo kết quả đầu ra và theo tỷ lệ phần trăm",
  'DC_TTG_CPCP_TLPT':
    "Đơn giá điều chỉnh, theo thời gian, theo chi phí cộng phí và theo tỷ lệ phần trăm",
  'DC_TTG_KQDR': "Đơn giá điều chỉnh, theo thời gian và theo kết quả đầu ra",
  'DC_TTG_KQDR_TLPT':
    "Đơn giá điều chỉnh, theo thời gian, theo kết quả đầu ra và theo tỷ lệ phần trăm",
  'DC_TTG_TLPT': "Đơn giá điều chỉnh, theo thời gian và theo tỷ lệ phần trăm",
  'DC_CPCP': "Đơn giá điều chỉnh và theo chi phí cộng phí",
  'DC_CPCP_KQDR':
    "Đơn giá điều chỉnh, theo chi phí cộng phí và theo kết quả đầu ra",
  'DC_CPCP_KQDR_TLPT':
    "Đơn giá điều chỉnh, theo chi phí cộng phí, theo kết quả đầu ra và theo tỷ lệ phần trăm",
  'DC_CPCP_TLPT':
    "Đơn giá điều chỉnh, theo chi phí cộng phí và theo tỷ lệ phần trăm",
  'DC_KQDR': "Đơn giá điều chỉnh và theo kết quả đầu ra",
  'DC_KQDR_TLPT':
    "Đơn giá điều chỉnh, theo kết quả đầu ra và theo tỷ lệ phần trăm",
  'DC_TLPT': "Đơn giá điều chỉnh và theo tỷ lệ phần trăm",
  'TTG': "Theo thời gian",
  'TTG_CPCP': "Theo thời gian và theo chi phí cộng phí",
  'TTG_CPCP_KQDR': "Theo thời gian, theo chi phí cộng phí và theo kết quả đầu ra",
  'TTG_CPCP_KQDR_TLPT':
    "Theo thời gian, theo chi phí cộng phí, theo kết quả đầu ra và theo tỷ lệ phần trăm",
  'TTG_CPCP_TLPT':
    "Theo thời gian, theo chi phí cộng phí và theo tỷ lệ phần trăm",
  'TTG_KQDR': "Theo thời gian và theo kết quả đầu ra",
  'TTG_KQDR_TLPT': "Theo thời gian, theo kết quả đầu ra và theo tỷ lệ phần trăm",
  'TTG_TLPT': "Theo thời gian và theo tỷ lệ phần trăm",
  'CPCP': "Theo chi phí cộng phí",
  'CPCP_KQDR': "Theo chi phí cộng phí và theo kết quả đầu ra",
  'CPCP_KQDR_TLPT':
    "Theo chi phí cộng phí, theo kết quả đầu ra và theo tỷ lệ phần trăm",
  'CPCP_TLPT': "Theo chi phí cộng phí và theo tỷ lệ phần trăm",
  'KQDR': "Theo kết quả đầu ra",
  'KQDR_TLPT': "Theo kết quả đầu ra và theo tỷ lệ phần trăm",
  'TLPT': "Theo tỷ lệ phần trăm",
  'KHAC': "Khác",
}

TBMT_STATUS_MAP = {
  'DXT': "Đang xét thầu",
  'DHTBMT': "Đã hủy TBMT",
  'DHT': "Đã huỷ thầu",
  'CNTTT': "Có nhà thầu trúng thầu",
  'KCNTT': "Không có nhà thầu trúng thầu",
  'VHH': "Tuyên bố vô hiệu quyết định về KQLCNT",
}

def map_quy_trinh_ap_dung(code):
    return PROCESS_APPLY_MAP.get(code, code)


def map_hinh_thuc_lua_chon_nha_thau(code):
    return BID_FORM_MAP.get(code, code)


def map_phuong_thuc_lua_chon_nha_thau(code):
    return BID_MODE_MAP.get(code, code)


def map_linh_vuc(code):
    return INVEST_FIELD_MAP.get(code, code)


def map_loai_hop_dong(code):
    return CONTRACT_TYPE_MAP.get(code, code)


def map_trang_thai_tbmt(code):
    clean_code = str(code).strip() if code else ""
    if not clean_code:
        return "Chưa đóng thầu"
    return TBMT_STATUS_MAP.get(clean_code, clean_code)


def get_trang_thai_dong_thau(thoi_diem_dong_thau, crawl_time):
    if not thoi_diem_dong_thau or not crawl_time:
        return ""

    close_time = js_date(thoi_diem_dong_thau)
    current_time = js_date(crawl_time)

    if not close_time or not current_time:
        return ""

    return (
        "Đã đóng thầu"
        if current_time >= close_time
        else "Chưa đóng thầu"
    )


def build_url_view_detail_tbmt(notify_id, process_apply):
    params = {
        "p_p_id": "egpportalcontractorselectionv2_WAR_egpportalcontractorselectionv2",
        "p_p_lifecycle": "0",
        "p_p_state": "normal",
        "p_p_mode": "view",
        "_egpportalcontractorselectionv2_WAR_egpportalcontractorselectionv2_render": "detail-v2",
        "type": "es-notify-contractor",
        "stepCode": "notify-contractor-step-1-tbmt",
        "id": notify_id,
        "notifyId": notify_id,
        "processApply": process_apply,
    }

    return (
        "https://muasamcong.mpi.gov.vn/web/guest/"
        "contractor-selection?"
        + urlencode(params)
    )


def build_url_view_detail_khlcnt(plan_id, plan_no):
    params = {
        "p_p_id": "egpportalcontractorselectionv2_WAR_egpportalcontractorselectionv2",
        "p_p_lifecycle": "0",
        "p_p_state": "normal",
        "p_p_mode": "view",
        "_egpportalcontractorselectionv2_WAR_egpportalcontractorselectionv2_render": "detail-v2",
        "type": "es-plan-project-p",
        "stepCode": "plan-step-1",
        "id": plan_id,
        "planNo": plan_no,
    }

    return (
        "https://muasamcong.mpi.gov.vn/web/guest/"
        "contractor-selection?"
        + urlencode(params)
    )


def build_detail_message(data):
    if not data:
        return "", ""

    gia_goi_thau = data.get("giaGoiThau")
    if gia_goi_thau:
        try:
            gia_goi_thau = f"{float(str(gia_goi_thau).replace(',', '')):,.0f}"
        except (ValueError, TypeError):
            gia_goi_thau = str(gia_goi_thau)
    else:
        gia_goi_thau = ""

    message_text = f"""<b>Mã TBMT:</b> {data.get('maTBMT', '')}
<b>Tên gói thầu:</b>
{data.get('tenGoiThau', '')}
<b>Giá gói thầu:</b> {gia_goi_thau}
<b>Chủ đầu tư:</b> {data.get('chuDauTu', '')}
<b>Mã KHLCNT:</b> {data.get('maKHLCNT', '')}
<b>Hình thức lựa chọn nhà thầu:</b> {map_hinh_thuc_lua_chon_nha_thau(data.get('hinhThucLuaChonNhaThau', ''))}
<b>Phương thức lựa chọn Nhà thầu:</b> {map_phuong_thuc_lua_chon_nha_thau(data.get('phuongThucLuaChonNhaThau', ''))}
<b>Loại hợp đồng:</b> {map_loai_hop_dong(data.get('loaiHopDong', ''))}
<b>Thời điểm đóng thầu:</b> {format_datetime(data.get('thoiDiemDongThau'))}
<b>Thời đăng tải gốc:</b> {format_datetime(data.get('ngayDangTaiGoc'))}
<b>Thời gian sửa TBMT:</b> {format_datetime(data.get('thoiGianSuaTBMT'))}"""

    url = build_url_view_detail_tbmt(data.get("id"), data.get("quyTrinhApDung"))
    return message_text.strip(), url


def js_date(value):
    if value is None:
        return None

    if isinstance(value, datetime):
        return value

    if isinstance(value, str):
        value = value.replace("Z", "+00:00")

        try:
            return datetime.fromisoformat(value)
        except ValueError:
            for fmt in ("%Y-%m-%d %H:%M:%S", "%Y-%m-%d"):
                try:
                    return datetime.strptime(value, fmt)
                except ValueError:
                    pass

    return None


def format_datetime(value):
    value = js_date(value)
    return value.strftime("%Y-%m-%d %H:%M:%S") if value else ""

def _clean_latex_and_text(text: str) -> str:
    """Tự động làm sạch mọi tàn dư LaTeX, công thức toán và HTML lạ."""
    if not text:
        return ""

    # 1. Bóc text thuần từ \text{...}
    text = re.sub(r"\\text\{([^}]+)\}", r"\1", text)

    # 2. Thay thế toàn bộ lệnh toán học LaTeX sang ký tự Unicode đẹp
    replacements = {
        r"\\ge\b": "≥",
        r"\\geq\b": "≥",
        r"\\le\b": "≤",
        r"\\leq\b": "≤",
        r"\\ne\b": "≠",
        r"\\neq\b": "≠",
        r"\\times\b": "×",
        r"\\approx\b": "≈",
        r"\\pm\b": "±",
        r"\\mp\b": "∓",
        r"\\infty\b": "∞",
    }
    for pattern, repl in replacements.items():
        text = re.sub(pattern, repl, text)

    # 3. Gỡ các bao đóng LaTeX: \(...\), \[...\], $...$
    text = re.sub(r"\\\((.*?)\\\)", r"\1", text)
    text = re.sub(r"\\\[(.*?)\\\]", r"\1", text)
    text = text.replace("$", "")

    # 4. Xóa dấu backslash (\) thừa phía trước các dấu ngoặc
    text = re.sub(r"\\([\\(\\)\[\]{}])", r"\1", text)

    return text


def _add_formatted_text(paragraph, text: str, is_bold_all: bool = False):
    """
    Bóc tách và định dạng:
    - Xử lý xuống dòng qua thẻ <br>, <br/>, <br />
    - Xử lý in đậm **...**, in nghiêng *...*, mã `...`
    """
    if not text:
        return

    # Làm sạch LaTeX trước
    text = _clean_latex_and_text(text)

    # Cắt dòng theo thẻ <br>
    sub_lines = re.split(r"(?i)<br\s*/?>", text)

    for line_idx, line in enumerate(sub_lines):
        line = line.strip()
        tokens = re.split(r"(\*\*.*?\*\*|\*.*?\*|`.*?`)", line)

        for token in tokens:
            if not token:
                continue

            if token.startswith("**") and token.endswith("**") and len(token) >= 4:
                run = paragraph.add_run(token[2:-2])
                run.bold = True
            elif token.startswith("*") and token.endswith("*") and len(token) >= 2:
                run = paragraph.add_run(token[1:-1])
                run.italic = True
                if is_bold_all:
                    run.bold = True
            elif token.startswith("`") and token.endswith("`") and len(token) >= 2:
                run = paragraph.add_run(token[1:-1])
                run.font.name = "Consolas"
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(180, 40, 40)
            else:
                run = paragraph.add_run(token)
                if is_bold_all:
                    run.bold = True

        # Nếu có thẻ <br>, thêm ngắt dòng trong ô/đoạn
        if line_idx < len(sub_lines) - 1:
            paragraph.add_run().add_break()


def _create_word_table(doc, table_data):
    """Tạo bảng chuẩn Word, viền Table Grid và tô màu header."""
    if not table_data:
        return

    num_rows = len(table_data)
    num_cols = max(len(row) for row in table_data)
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = "Table Grid"

    for r_idx, row in enumerate(table_data):
        is_header = (r_idx == 0)
        for c_idx in range(num_cols):
            cell_text = row[c_idx] if c_idx < len(row) else ""
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            p = cell.paragraphs[0]

            _add_formatted_text(p, cell_text, is_bold_all=is_header)

            if is_header:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F4F8"/>')
                cell._tc.get_or_add_tcPr().append(shading)


def save_markdown_to_docx(markdown_text: str, output_path: str, title: str = ""):
    """Tự động phân tích toàn bộ cấu trúc Markdown xuất sang file Word."""
    doc = Document()

    if title:
        main_title = doc.add_heading(title, level=0)
        main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    cleaned_text = markdown_text.strip()
    if cleaned_text.startswith("```"):
        cleaned_text = re.sub(r"^```[a-zA-Z]*\n?", "", cleaned_text)
        cleaned_text = re.sub(r"\n?```$", "", cleaned_text)

    lines = cleaned_text.splitlines()
    table_buffer = []
    in_table = False
    in_code_block = False
    code_buffer = []

    for line in lines:
        stripped = line.strip()

        # 1. Code Block
        if stripped.startswith("```"):
            if in_code_block:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.3)
                run = p.add_run("\n".join(code_buffer))
                run.font.name = "Consolas"
                run.font.size = Pt(9.5)
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_buffer.append(line)
            continue

        # 2. Xử lý Bảng
        if stripped.startswith("|") and stripped.endswith("|"):
            if re.match(r"^\|[\s\-:|]+\|$", stripped):
                continue
            cols = [col.replace(r"\|", "|").strip() for col in re.split(r"(?<!\\)\|", stripped)[1:-1]]
            table_buffer.append(cols)
            in_table = True
            continue
        else:
            if in_table:
                _create_word_table(doc, table_buffer)
                table_buffer = []
                in_table = False

        if not stripped:
            continue

        # 3. Kẻ ngang ---
        if re.match(r"^(?:-{3,}|\*{3,}|_{3,})$", stripped):
            p = doc.add_paragraph()
            run = p.add_run("―" * 40)
            run.font.color.rgb = RGBColor(160, 160, 160)
            continue

        # 4. Blockquote >
        if stripped.startswith(">"):
            quote_text = re.sub(r"^>\s?", "", stripped)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            run = p.add_run("“ ")
            run.italic = True
            run.font.color.rgb = RGBColor(90, 90, 90)
            _add_formatted_text(p, quote_text)
            continue

        # 5. Heading (#, ##, ###)
        heading_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading_match:
            level = min(len(heading_match.group(1)), 3)
            doc.add_heading(_clean_latex_and_text(heading_match.group(2)), level=level)
            continue

        # 6. Danh sách Bullet (- hoặc *)
        if re.match(r"^[-*+]\s+", stripped):
            content = re.sub(r"^[-*+]\s+", "", stripped)
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_text(p, content)
            continue

        # 7. Danh sách có số (1., 2.)
        if re.match(r"^\d+[\.\)]\s+", stripped):
            content = re.sub(r"^\d+[\.\)]\s+", "", stripped)
            p = doc.add_paragraph(style="List Number")
            _add_formatted_text(p, content)
            continue

        # 8. Văn bản thường
        p = doc.add_paragraph()
        _add_formatted_text(p, stripped)

    if in_table and table_buffer:
        _create_word_table(doc, table_buffer)

    doc.save(output_path)